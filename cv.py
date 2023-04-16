import os
import json
import subprocess
from docx import Document
from time import sleep
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell
from docx.text.paragraph import Paragraph, Parented
from docx.shared import Pt, Inches
from collections.abc import Callable
from utils import (
    insertHR,
    format_date_range,
    add_hyperlink,
    format_year_range,
    parse_date,
    indent_table,
    add_page_number,
    BLUE,
    DARK_BLUE,
    GRAY,
    LIGHT_GRAY,
    BLACK
)


class CV:
    """ Curriculum Vitae class """
    CV_KEY = 'cv'
    WORKS_KEY = 'works'

    def __init__(self, font: str = 'Lato', reverse_format: bool = False) -> None:
        self.doc = Document()
        self.font = font
        self.font_size = 10.5
        self.tab_size = Inches(0.3)
        self.date_col_width = Inches(0.95)
        self.item_col_width = Inches(5.25)
        self.reverse_format = reverse_format
        for style_name in ['Normal'] + [f'Heading {i}' for i in range(1, 6)]:
            style = self.doc.styles[style_name]
            style.font.name = self.font
            style.font.size = Pt(self.font_size)
            style.font.color.rgb = BLACK
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)

    def write(self, output_file: str, open_file: bool = True) -> None:
        self.doc.save(output_file)
        if open_file:
            cmd = """osascript -e 'tell application "Microsoft Word" to close windows'"""
            os.system(cmd)
            sleep(1)
            subprocess.run(['open', output_file])

    def load_data(self, cv_path: str, works_path: str) -> None:
        data = {}
        for json_file, key in zip([cv_path, works_path], [self.CV_KEY, self.WORKS_KEY]):
            with open(json_file, 'r') as f:
                data[key] = json.load(f)
        self.data = data

    def compile(self) -> None:
        self.__apply_formatting()
        self.__parse_basics()
        self.__parse_education()
        self.__parse_experience()
        self.__parse_publications()
        self.__parse_awards()
        self.__parse_skills()
        self.__parse_works()

    def __new_section(self, name: str) -> Paragraph:
        self.__insert_break(2)
        header = self.doc.add_heading(name)
        header.runs[0].font.name = self.font
        insertHR(header)
        return header

    def __new_subsection(self, name: str) -> Paragraph:
        self.__insert_break()
        subheader = self.doc.add_heading(name, level=2)
        font = subheader.runs[0].font
        font.name = self.font
        font.color.rgb = DARK_BLUE
        self.__insert_break()
        return subheader

    def __insert_break(self, n_units: int | float = 1, parent: Parented | None = None):
        obj = parent or self.doc
        gap = Pt(7 * n_units)
        p = obj.add_paragraph(" ")
        p.runs[0].font.size = gap
        p.paragraph_format.line_spacing = gap

    def __make_entry_table(self, parent: object, items: list, handler: Callable, date_getter: Callable) -> None:
        tbl = parent.add_table(len(items), 2)
        last_date = None
        for i, item in enumerate(items):
            date_cell, item_cell = tbl.cell(i, int(self.reverse_format)), tbl.cell(i, 1-int(self.reverse_format))
            date_cell.width, item_cell.width = self.date_col_width, self.item_col_width
            date = date_getter(item)
            if date != last_date:
                date_cell.paragraphs[0].add_run(date)
            last_date = date
            last_element = handler(cell=item_cell, item=item)
            self.__insert_break(0.5, parent=last_element or date_cell)

    def __parse_basics(self) -> None:
        self.__parse_personal_info()
        self.__parse_interests()

    def __parse_personal_info(self) -> None:
        basics = self.data[self.CV_KEY]['basics']

        p = self.doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name = p.add_run(basics['name'].upper())
        name.bold = True

        labels = p.add_run("\n{} | Curriculum Vitae".format(" + ".join([l.capitalize() for l in basics['labels']])))
        labels.bold = True

        address_key = p.add_run("\nAddress: ")
        address_key.bold = True

        p.add_run("{}, {}, {} — {}. ".format(*[basics['location'][x]
                                               for x in ['address', 'city', 'region', 'countryCode',]]))

        phone_key = p.add_run("Phone: ")
        phone_key.bold = True

        _phone = basics['phone']
        p.add_run(_phone)

        _url = basics['profiles'][-1]['url']
        p.add_run(f"\n")
        add_hyperlink(p, _url, _url)
        p.add_run(f" | ")
        add_hyperlink(p, basics['email'], basics['email'])

    def __apply_formatting(self) -> None:
        basics = self.data[self.CV_KEY]['basics']
        self.doc.settings.odd_and_even_pages_header_footer = True
        self.doc.sections[0].different_first_page_header_footer = True
        for i, attr in enumerate(['header', 'even_page_header']):
            p = getattr(self.doc.sections[0], attr).paragraphs[0]
            tab = '\t\t'
            header = p.add_run(f"{[tab, ''][i]}{basics['name'].upper()}")
            header.bold = True
            suffix = p.add_run(' | Curriculum Vitae')
            for x in [header, suffix]:
                x.font.color.rgb = LIGHT_GRAY
        self.doc.sections[0].first_page_header.paragraphs[0].text = ''
        for i, attr in enumerate(['footer', 'even_page_footer']):
            p = getattr(self.doc.sections[0], attr).paragraphs[0]
            add_page_number(p.add_run())
            p.runs[0].font.color.rgb = LIGHT_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __parse_interests(self) -> None:
        self.__insert_break(2)
        interests = [x.lower() if x[1].islower() else x for x in self.data[self.CV_KEY]['basics']['interests']]
        interests.sort()
        p = self.doc.add_paragraph("Interests: ")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].bold = True
        f = p.paragraph_format
        f.left_indent = self.tab_size * 2
        f.right_indent = self.tab_size * 2
        keywords = p.add_run(f'{" • ".join(interests)}.')
        keywords.font.color.rgb = GRAY

    def __parse_education(self) -> None:
        education = self.data[self.CV_KEY].pop('education', None)
        if not education:
            return
        self.__new_section("EDUCATION")
        degrees = education.pop('degrees', None)
        if degrees:
            self.__new_subsection("Degrees")
            for degree in degrees:
                p = self.doc.add_paragraph()
                name = p.add_run(degree['name'])
                name.bold = True

                p.add_run(f", {degree['major']}.")

                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = self.tab_size

                institution = p.add_run(f"{degree['institution']}")
                institution.bold = True

                p.add_run(f". {degree['city']}. {degree['country']}. {format_year_range(*degree['date'])}.")
                minors = degree['minors']
                if minors:
                    minors_label = p.add_run("\nMinor fields: ")
                    minors_label.font.bold = True
                    p.add_run(f"{', '.join(degree['minors'])}.")
                highlights = degree['highlights']
                if highlights:
                    lb = p.add_run(f"\nHighlights: ")
                    lb.bold = True
                    hls = p.add_run(f"{', '.join(highlights)}.")
                    hls.italic = True

        other_ed = education.pop('other', None)
        if other_ed:
            self.__new_subsection("Other")
            for ed in other_ed:
                p = self.doc.add_paragraph("")
                p.paragraph_format.left_indent = self.tab_size
                p.paragraph_format.first_line_indent = -self.tab_size

                name = p.add_run(ed['name'])
                name.bold = True
                p.add_run(f" ({ed['type']}). {ed['institution']}. {ed['location']}. {format_date_range(*ed['date'])}.")

    def __parse_experience(self) -> None:
        self.__parse_jobs()
        self.__parse_lectures()
        self.__parse_workshops()
        self.__parse_residencies()

    def __parse_jobs(self) -> None:
        def handler(cell: _Cell, item: dict) -> _Cell:
            position = item
            p = cell.paragraphs[0]
            name = p.add_run(position['name'])
            name.bold = True

            p.add_run(f". {position['workplace']}. {position['city']}. {position['country']}.")

            _courses = position.pop('courses', None)
            if _courses:
                course_tbl = cell.add_table(len(_courses), 2)
                course_label_cell = course_tbl.cell(0, 0)
                course_label_cell.paragraphs[0].add_run("Courses:")

                for j, course in enumerate(_courses):
                    course_cell = course_tbl.cell(j, 1)
                    course_cell.width = Inches(20)
                    cp = course_cell.paragraphs[0]
                    course_name = cp.add_run(course['name'])
                    course_name.font.bold = True
                    terms = cp.add_run(", {}.".format(course['terms']))
                    terms.font.color.rgb = GRAY
                return course_cell

        def date_getter(item: dict) -> str:
            return format_year_range(*item['date'])

        work = self.data[self.CV_KEY]['work']
        self.__new_section("WORK EXPERIENCE")
        academic_work = work.pop('academic', None)
        other_work = work.pop('other positions', None)
        experience_sections = [x for x in [("Teaching experience", academic_work),
                                           ("Other positions", other_work)] if x[1] is not None]
        for label, positions in experience_sections:
            self.__new_subsection(label)
            positions.sort(key=lambda x: 10000 if x['date'][1] ==
                           True else (x['date'][0] if x['date'][1] == False else x['date'][1]), reverse=True)
            self.__make_entry_table(self.doc, positions, handler, date_getter)

    def __parse_lectures(self) -> None:
        lectures = self.data[self.CV_KEY]['work'].pop('lectures', None)
        if not lectures:
            return
        self.__new_subsection("Guest lectures")
        lectures.sort(key=lambda x: sorted(x['events'], key=lambda y: y['date'], reverse=True)[0]['date'], reverse=True)
        for lecture in lectures:
            p = self.doc.add_paragraph()
            name = p.add_run(lecture['name'])
            name.bold = True
            name.italic = True

            for event in lecture['events']:
                p = self.doc.add_paragraph("@ ")
                p.runs[0].font.color.rgb = BLUE
                p.paragraph_format.left_indent = self.tab_size
                event_name = p.add_run(event['name'])
                event_name.font.italic = True
                year, month, day = parse_date(event['date'])
                p.add_run(f". { event['venue']}. {event['city']}. {event['country']}. {f'{month} {day}, {year}'}.")

    def __parse_workshops(self) -> None:
        workshops = self.data[self.CV_KEY]['work'].pop('workshops', None)
        if not workshops:
            return
        self.__new_subsection("Workshops")
        workshops.sort(key=lambda x: sorted(x['events'], key=lambda y: y['date'], reverse=True)[0]['date'], reverse=True)

        for workshop in workshops:
            p = self.doc.add_paragraph()
            name = p.add_run(workshop['name'])
            name.bold = True
            name.italic = True

            for event in workshop['events']:
                p = self.doc.add_paragraph("@ ")
                p.runs[0].font.color.rgb = BLUE
                p.paragraph_format.left_indent = self.tab_size
                institution = p.add_run(event['institution'])
                institution.italic = True
                year, month, day = parse_date(event['date'])

                p.add_run(
                    f". {event['numSessions']} sessions ({event['totalHours']} hours total). {event['city']}. {event['country']}. {month} {day}, {year}.")

    def __parse_residencies(self) -> None:
        residencies = self.data[self.CV_KEY]['work'].pop('residencies', None)
        if not residencies:
            return

        def date_getter(item: dict) -> str:
            return parse_date(item['date'])[0]

        def handler(cell: _Cell, item: dict) -> None:
            residency = item
            p = cell.paragraphs[0]
            role = p.add_run(residency['role'])
            role.font.italic = True
            at = p.add_run(" @ ")
            at.font.color.rgb = BLUE
            event = p.add_run(residency['event'])
            event.font.bold = True

            _institution = residency['institution']
            _end = residency['end']
            date_range = format_date_range(residency['date'], _end)
            p.add_run(f". {_institution}. {date_range}.")

            p = cell.add_paragraph()
            p.paragraph_format.left_indent = self.tab_size
            label = p.add_run("Activities: ")
            label.font.bold = True
            _activities = "{}.".format(", ".join(residency['activities']))
            activities = p.add_run(_activities)
            activities.font.italic = True

        self.__new_subsection("Residencies")
        residencies.sort(key=lambda x: x['date'], reverse=True)
        self.__make_entry_table(self.doc, residencies, handler, date_getter)

    def __parse_awards(self) -> None:
        self.__new_section("AWARDS")
        awards = []
        commissions = []
        for work in self.data[self.WORKS_KEY]:
            work_awards = work['awards']
            commission = work['commission']
            if work_awards:
                for award in work_awards:
                    awards.append(award)
            if commission:
                commissions.append(work)

        self.___parse_awards(awards, "Artistic awards")
        self.___parse_commissions(commissions)
        self.___parse_awards(self.data[self.CV_KEY]['awards'].pop('academic', None), "Academic awards")

    def ___parse_commissions(self, commissions) -> None:
        if not commissions:
            return

        def handler(cell: _Cell, item: dict) -> None:
            commission = item
            p = cell.paragraphs[0]
            name = p.add_run(commission['name'])
            name.font.bold = True

            subtitle = p.add_run(" {}.".format(commission['subtitle']))
            subtitle.font.italic = True

            p.add_run(" {}.".format(commission['commission']))

        def date_getter(item: dict) -> str:
            return str(item['year'])

        self.__new_subsection("Commissions")
        commissions.sort(key=lambda x: x['year'], reverse=True)
        self.__make_entry_table(self.doc, commissions, handler, date_getter)

    def ___parse_awards(self, awards: dict, label: str) -> None:
        if not awards:
            return

        def handler(cell: _Cell, item: dict) -> None:
            award = item
            p = cell.paragraphs[0]
            name = p.add_run(award['name'])
            name.font.bold = True

            p.add_run(". {}. {}.".format(*[award[x] for x in ['institution', 'country']]))

        def date_getter(item: dict) -> str:
            return str(item['date'])

        self.__new_subsection(label)
        awards.sort(key=lambda x: x['date'], reverse=True)
        self.__make_entry_table(self.doc, awards, handler, date_getter)

    def __parse_publications(self) -> None:
        publications = self.data[self.CV_KEY]['work'].pop('publications', None)
        if publications:
            def handler(cell: _Cell, item: dict) -> None:
                pub = item
                p = cell.paragraphs[0]
                p.add_run(f"{pub['author']} ({pub['date']}). ")
                name = p.add_run(pub['name'])
                name.font.bold = True

                publisher = p.add_run(f". {pub['publisher']}")
                publisher.font.italic = True

                p.add_run(f", ({pub['edition']}), {'-'.join([str(x) for x in pub['pages']])}. ")
                add_hyperlink(p, pub['doi'], pub['doi'])

            def date_getter(item: dict) -> str:
                return str(item['date'])

            def recording_handler(cell: _Cell, item: dict) -> None:
                rec = item
                p = cell.paragraphs[0]
                album = p.add_run(f"{rec['album']}. ")
                album.italic = True

                track = p.add_run(f"{rec['track']}. ")
                track.bold = True

                p.add_run(f"{rec['recordLabel']}. ")

                performers = rec.pop('performers', None)
                if not performers:
                    return
                num_performers = len(performers)
                for i, performer in enumerate(performers):
                    p.add_run(f"{performer['name']} ")
                    role = p.add_run(f"({performer['role']}){', ' if i < num_performers - 1 else '.'}")
                    role.italic = True

            def recording_date_getter(item: str) -> str:
                return str(item['year'])

            self.__new_section("PUBLICATIONS")
            articles, scores, recordings = publications['articles'], publications['scores'], publications['recordings']
            for items, label in [(articles, 'Peer-reviewed articles'), (scores, 'Scores'), (recordings, 'Recordings')]:
                self.__new_subsection(label)
                items.sort(key=lambda x: x['date'] if 'date' in x else x['year'], reverse=True)
                handle_func = handler if label != 'Recordings' else recording_handler
                date_func = date_getter if label != 'Recordings' else recording_date_getter
                self.__make_entry_table(self.doc, items, handle_func, date_func)

        software_list = self.data[self.CV_KEY]['work'].pop('software', None)
        if software_list:
            def handler(cell: _Cell, item: dict) -> None:
                software = item
                p = cell.paragraphs[0]
                name = p.add_run(software['name'])
                name.font.bold = True

                url = software['url']
                p.add_run(" (")
                add_hyperlink(p, url, url)
                p.add_run(")")

                p = cell.add_paragraph("Keywords: ")
                p.runs[0].italic = p.runs[0].bold = True
                p.paragraph_format.left_indent = self.tab_size
                keywords = p.add_run(f'{", ".join(software["keywords"])}.')
                keywords.italic = True

                p = cell.add_paragraph("Description: ")
                p.runs[0].italic = p.runs[0].bold = True
                p.paragraph_format.left_indent = self.tab_size
                descr = p.add_run(f"{software['description']}")
                descr.font.italic = True

            def date_getter(item: dict) -> str:
                return str(item['year'])

            self.__new_subsection("Software")
            software_list.sort(key=lambda x: x['year'], reverse=True)
            self.__make_entry_table(self.doc, software_list, handler, date_getter)

    def __parse_skills(self) -> None:
        self.__new_section("SKILLS")
        skills_dict = self.data[self.CV_KEY].pop('skills', None)
        if not skills_dict:
            return
        divs = 3
        for skill_key in skills_dict:
            self.__new_subsection(skill_key.capitalize())
            skills = skills_dict[skill_key]
            skills.sort(key=lambda x: x['level'])
            tbl = self.doc.add_table(round(len(skills) / divs), divs)
            indent_table(tbl, 350)
            for i, skill in enumerate(skills):
                row = i // divs
                col = i % divs
                cell = tbl.cell(row, col)

                p = cell.paragraphs[0]

                name = p.add_run(skill['name'])
                name.font.bold = True

                keywords = p.add_run(f' ({", ".join(skill["keywords"])})')
                keywords.font.italic = True

                gap = Pt(5)

                p = cell.add_paragraph(" ")
                p.runs[0].font.size = gap
                p.paragraph_format.line_spacing = gap
            self.__insert_break()

    def __parse_works(self) -> None:
        works = self.data[self.WORKS_KEY]
        if not works:
            return
        self.doc.add_page_break()
        self.__new_section("LIST OF WORKS")
        works.sort(key=lambda x: x['year'], reverse=True)
        last_date = None
        for i, work in enumerate(works):
            date = str(work['year'])
            if date != last_date:
                self.__new_subsection(f"{date}")
            last_date = date
            p = self.doc.add_paragraph()
            name = p.add_run(work['name'])
            name.font.bold = True

            p.add_run(f" ({date}) ")
            subtitle = p.add_run(f"{work['subtitle']}. ")
            subtitle.italic = True

            p.add_run(f"{work['duration']}'")

            _commission = work['commission']
            if _commission:
                self.__insert_break(0.5)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = self.tab_size
                commission = p.add_run(f'{_commission}.')
                commission.italic = True
                commission.font.color.rgb = GRAY

            performances = work['performances']
            if performances:
                performances.sort(key=lambda x: x['date'], reverse=True)
                num_perf = len(performances)
                self.__insert_break(0.5)
                p = self.doc.add_paragraph("Performances")
                label = p.runs[0]
                label.italic = True
                label.bold = True
                label.font.color.rgb = DARK_BLUE
                p.paragraph_format.left_indent = self.tab_size
                self.__insert_break(0.5)
                for i, performance in enumerate(performances):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = self.tab_size

                    event = p.add_run(performance['event'])
                    event.font.bold = True

                    if i == num_perf - 1:
                        p.add_run(" (world premiere)")

                    at = p.add_run(" @ ")
                    at.font.color.rgb = BLUE

                    year, month, day = parse_date(performance['date'])
                    p.add_run(
                        f"{performance['venue']}. {performance['city']}. {performance['country']}. {month} {day}, {year}. ")

                    performers = performance.pop('performers', None)
                    if performers:
                        p = self.doc.add_paragraph("Performed by ")
                        p.paragraph_format.left_indent = self.tab_size * 2
                        num_performers = len(performers)
                        for i, performer in enumerate(performers):
                            p.add_run(performer['name'])
                            role = p.add_run(
                                f" ({performer['role']}){'.' if i == num_performers - 1 else (', and ' if i == num_performers - 2 else ', ')}")
                            role.italic = True
                    self.__insert_break(0.5)
