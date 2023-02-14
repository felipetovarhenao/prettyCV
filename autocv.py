from docx import Document
import json
import subprocess
from time import sleep
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import os
from utils import (
    insertHR,
    format_date_range,
    add_hyperlink,
    format_year_range,
    parse_date,
    indent_table,
    add_page_number,
    BLUE,
    BLACK
)


class CV:
    CV_KEY = 'cv'
    WORKS_KEY = 'works'

    def __init__(self, font='Nunito') -> None:
        self.doc = Document()
        self.font = font
        self.font_size = 10.5
        self.tab_size = Inches(0.25)
        self.date_col_width = Inches(0.85)
        self.item_col_width = Inches(5)
        for style_name in ['Normal', *[f'Heading {i}' for i in range(1, 6)]]:
            style = self.doc.styles[style_name]
            style.font.name = self.font
            style.font.size = Pt(self.font_size)
            style.font.color.rgb = BLACK
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)
        add_page_number(self.doc.sections[0].footer.paragraphs[0].add_run())
        self.doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def write(self, output_file):
        self.doc.save(output_file)
        cmd = """osascript -e 'tell application "Microsoft Word" to close windows'"""
        os.system(cmd)
        sleep(1)
        subprocess.run(['open', output_file])

    def load_data(self, cv_path, works_path):
        data = {}
        for json_file, key in zip([cv_path, works_path], [self.CV_KEY, self.WORKS_KEY]):
            with open(json_file, 'r') as f:
                data[key] = json.load(f)
        self.data = data

    def compile(self):
        self._add_main_header()
        self._add_education()
        self._add_work_experience()
        self._add_awards()
        self._add_publications()
        self._add_skills()
        self.doc.add_page_break()
        self._add_works()

    def _add_main_header(self):
        basics = self.data[self.CV_KEY]['basics']

        p = self.doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name = p.add_run(basics['name'].upper())
        name.bold = True

        labels = p.add_run("\n{} | Curriculum Vitae".format(" + ".join([l.capitalize() for l in basics['labels']])))
        labels.bold = True

        info_key = p.add_run("\nAddress: ")
        info_key.bold = True

        p.add_run("{}, {}, {} — {}. ".format(*[basics['location'][x]
                                               for x in ['address', 'city', 'region', 'countryCode',]]))

        info_key = p.add_run("Phone: ")
        info_key.bold = True

        _phone = basics['phone']
        p.add_run(_phone)

        _email = basics['email']
        _url = basics['profiles'][-1]['url']
        p.add_run(f"\n")
        add_hyperlink(p, _url, _url)
        p.add_run(f" | ")
        add_hyperlink(p, _email, f"mailto:{_email}")
        self._add_break(2)

    def _new_section(self, name):
        header = self.doc.add_heading(name)
        insertHR(header)
        header.runs[0].font.name = self.font
        self._add_break()
        return header

    def _new_subsection(self, name):
        subheader = self.doc.add_heading(name, level=2)
        subheader.runs[0].font.name = self.font
        self._add_break()
        return subheader

    def _add_break(self, size=1):
        gap = Pt(size*7)
        p = self.doc.add_paragraph(" ")
        p.runs[0].font.size = gap
        p.paragraph_format.line_spacing = gap

    def _add_education(self):
        self._new_section("EDUCATION")
        education = self.data[self.CV_KEY]['education']
        degrees = education['degrees']
        for _degree in degrees:
            p = self.doc.add_paragraph("")
            _name = _degree['name']
            name = p.add_run(_name)
            name.font.bold = True

            _major = _degree['major']
            p.add_run(f', {_major}.')

            p = self.doc.add_paragraph("")
            p.paragraph_format.left_indent = self.tab_size

            _institution = _degree['institution']
            institution = p.add_run(f"{_institution}")
            institution.font.bold = True

            _city = _degree['city']
            _country = _degree['country']
            _dates = " - ".join([str(x) for x in _degree['date']])
            p.add_run(f". {_city}. {_country}. {_dates}.")
            minors = _degree['minors']
            if minors:
                minors_label = p.add_run("\nMinor fields: ")
                minors_label.font.bold = True
                p.add_run(f"{', '.join(_degree['minors'])}.")
        other_ed = education['other']
        self._add_break()

        for ed in other_ed:
            p = self.doc.add_paragraph("")
            p.paragraph_format.left_indent = self.tab_size
            p.paragraph_format.first_line_indent = -self.tab_size

            _name = ed['name']
            _url = ed['url']
            if _url:
                name = add_hyperlink(p, _name, _url)
            else:
                name = p.add_run({_name})
                name.font.bold = True

            _type = ed['type']
            _institution = ed['institution']
            _location = ed['location']
            _date = format_date_range(*ed['date'])

            p.add_run(f" ({_type}). {_location}. {_date}.")

        self._add_break(2)

    def _add_work_experience(self):
        work = self.data[self.CV_KEY]['work']
        self._new_section("WORK EXPERIENCE")
        experience_list = [work['academic'], work['other positions']]
        experience_labels = ["Teaching experience", "Other positions"]
        for positions, label in zip(experience_list, experience_labels):
            self._new_subsection(label)
            positions.sort(key=lambda x: 10000 if x['date'][1] ==
                           True else (x['date'][0] if x['date'][1] == False else x['date'][1]), reverse=True)
            tbl = self.doc.add_table(len(positions), 2)
            last_date = None
            for i, position in enumerate(positions):
                date_cell, position_cell = tbl.cell(i, 0), tbl.cell(i, 1)
                date_cell.width = self.date_col_width

                position_cell.width = self.item_col_width

                p = date_cell.paragraphs[0]

                date_val = format_year_range(*position['date'])
                if date_val != last_date:
                    p.add_run(date_val)
                last_date = date_val

                p = position_cell.paragraphs[0]
                name = p.add_run(position['name'])
                name.font.bold = True

                _workplace = position['workplace']
                _city = position['city']
                _country = position['country']

                p.add_run(f". {_workplace}. {_city}. {_country}.")

                _courses = position.pop('courses', None)
                if _courses:
                    num_courses = len(_courses)
                    course_tbl = position_cell.add_table(num_courses, 2)
                    course_tbl.autofit = False
                    course_tbl.allow_autofit = False
                    course_label_cell = course_tbl.cell(0, 0)
                    course_label_cell.paragraphs[0].add_run("Courses:")

                    for j, course in enumerate(_courses):
                        course_cell = course_tbl.cell(j, 1)
                        course_cell.width = Inches(20)
                        cp = course_cell.paragraphs[0]
                        course_name = cp.add_run(course['name'])
                        course_name.font.bold = True
                        cp.add_run(", {}.".format(course['terms']))
                    cp.add_run("\n")
        self._add_break(2)
        self._add_lectures()
        self._add_break(2)
        self.add_workshops()
        self._add_break(2)
        self.add_residencies()
        self._add_break(2)

    def _add_lectures(self):
        self._new_subsection("Guest lectures")
        lectures = self.data[self.CV_KEY]['work']['lectures']
        lectures.sort(key=lambda x: sorted(x['events'], key=lambda y: y['date'], reverse=True)[0]['date'], reverse=True)
        for lecture in lectures:
            p = self.doc.add_paragraph("")
            name = p.add_run(lecture['name'])
            name.font.bold = True
            name.font.italic = True

            for event in lecture['events']:
                p = self.doc.add_paragraph("@ ")
                p.runs[0].font.color.rgb = BLUE
                p.paragraph_format.left_indent = self.tab_size
                event_name = p.add_run(event['name'])
                event_name.font.italic = True
                _date = parse_date(event['date'])
                _date = f"{_date[1]} {_date[2]}, {_date[0]}"
                _venue = event['venue']
                _city = event['city']
                _country = event['country']

                p.add_run(f". {_venue}. {_city}. {_country}. {_date}.")

    def add_workshops(self):
        self._new_subsection("Workshops")
        workshops = self.data[self.CV_KEY]['work']['workshops']
        workshops.sort(key=lambda x: sorted(x['events'], key=lambda y: y['date'], reverse=True)[0]['date'], reverse=True)

        for workshop in self.data[self.CV_KEY]['work']['workshops']:
            p = self.doc.add_paragraph("")
            name = p.add_run(workshop['name'])
            name.font.bold = True
            name.font.italic = True

            for event in workshop['events']:
                p = self.doc.add_paragraph("@ ")
                p.runs[0].font.color.rgb = BLUE
                p.paragraph_format.left_indent = self.tab_size
                institution = p.add_run(event['institution'])
                institution.font.italic = True
                _date = parse_date(event['date'])
                _date = f"{_date[1]} {_date[2]}, {_date[0]}"
                _city = event['city']
                _country = event['country']
                _sessions = event['numSessions']
                _hours = event['totalHours']

                p.add_run(f". {_sessions} sessions ({_hours} hours total). {_city}. {_country}. {_date}.")

    def add_residencies(self):
        self._new_subsection("Residencies")
        residencies = self.data[self.CV_KEY]['work']['residencies']
        residencies.sort(key=lambda x: x['date'], reverse=True)
        tbl = self.doc.add_table(len(residencies), 2)
        for i, residency in enumerate(residencies):
            left, right = tbl.cell(i, 0), tbl.cell(i, 1)
            left.width, right.width = self.date_col_width, self.item_col_width

            _date = residency['date']

            p = left.paragraphs[0]
            p.add_run(parse_date(_date)[0])

            p = right.paragraphs[0]
            role = p.add_run(residency['role'])
            role.font.italic = True
            at = p.add_run(" @ ")
            at.font.color.rgb = BLUE
            event = p.add_run(residency['event'])
            event.font.bold = True

            _institution = residency['institution']
            _end = residency['end']
            date_range = format_date_range(_date, _end)
            p.add_run(f". {_institution}. {date_range}.")

            p = right.add_paragraph()
            p.paragraph_format.left_indent = self.tab_size
            label = p.add_run("Activities: ")
            label.font.bold = True
            _activities = "{}.".format(", ".join(residency['activities']))
            activities = p.add_run(_activities)
            activities.font.italic = True

    def _add_awards(self):
        self._new_section("AWARDS")
        self._new_subsection("Composition awards")
        awards = []
        commissions = []
        for work in self.data[self.WORKS_KEY]:
            work_awards = work['awards']
            commission = work['commission']
            if work_awards:
                for award in work_awards:
                    awards.append(award)
            if commission:
                commissions.append(work)

        awards.sort(key=lambda x: x['date'], reverse=True)
        tbl = self.doc.add_table(len(awards), 2)
        last_date = None
        for i, award in enumerate(awards):
            left, right = tbl.cell(i, 0), tbl.cell(i, 1)
            left.width, right.width = self.date_col_width, self.item_col_width

            p = left.paragraphs[0]
            date = str(award['date'])
            if date != last_date:
                p.add_run(date)
            last_date = date

            p = right.paragraphs[0]
            name = p.add_run(award['name'])
            name.font.bold = True

            p.add_run(". {}. {}.".format(*[award[x] for x in ['institution', 'country']]))

        self._add_break(2)
        self._new_subsection("Commissions")
        commissions.sort(key=lambda x: x['year'], reverse=True)
        tbl = self.doc.add_table(len(commissions), 2)
        last_date = None
        for i, commission in enumerate(commissions):
            left, right = tbl.cell(i, 0), tbl.cell(i, 1)
            left.width, right.width = self.date_col_width, self.item_col_width

            date = str(commission['year'])
            if date != last_date:
                p = left.paragraphs[0]
                p.add_run(date)
            last_date = date

            p = right.paragraphs[0]
            name = p.add_run(commission['name'])
            name.font.bold = True

            subtitle = p.add_run(" {}.".format(commission['subtitle']))
            subtitle.font.italic = True

            p.add_run(" {}.".format(commission['commission']))

        self._add_break(2)
        self._new_subsection("Academic awards")
        awards = self.data[self.CV_KEY]['awards']['academic']
        awards.sort(key=lambda x: x['date'], reverse=True)
        tbl = self.doc.add_table(len(awards), 2)
        last_date = None
        for i, award in enumerate(awards):
            left, right = tbl.cell(i, 0), tbl.cell(i, 1)
            left.width, right.width = self.date_col_width, self.item_col_width
            p = left.paragraphs[0]

            date = str(award['date'])
            if date != last_date:
                p.add_run(date)
            last_date = date

            p = right.paragraphs[0]
            name = p.add_run(award['name'])
            name.font.bold = True

            p.add_run(". {}. {}.".format(*[award[x] for x in ['institution', 'country']]))
        self._add_break(2)

    def _add_publications(self):
        self._new_section("PUBLICATIONS")
        publications = self.data[self.CV_KEY]['work']['publications']
        articles, scores = publications[:1], publications[1:]
        for items, label in zip([articles, scores], ['Peer-reviewed articles', 'Scores']):
            self._new_subsection(label)
            items.sort(key=lambda x: x['date'])
            tbl = self.doc.add_table(len(items), 2)
            last_date = None
            for i, pub in enumerate(items):
                left, right = tbl.cell(i, 0), tbl.cell(i, 1)
                left.width, right.width = self.date_col_width, self.item_col_width

                p = left.paragraphs[0]
                date = str(pub['date'])
                if date != last_date:
                    p.add_run(date)
                last_date = date

                p = right.paragraphs[0]
                p.add_run(f"{pub['author']} ({pub['date']}). ")
                name = p.add_run(pub['name'])
                name.font.bold = True

                publisher = p.add_run(f". {pub['publisher']}")
                publisher.font.italic = True

                p.add_run(f", ({pub['edition']}), {'-'.join([str(x) for x in pub['pages']])}. ")
                add_hyperlink(p, pub['doi'], pub['doi'])

        self._add_break(2)
        self._new_subsection("Software contributions")
        software_list = self.data[self.CV_KEY]['work']['software']
        software_list.sort(key=lambda x: x['year'], reverse=True)
        tbl = self.doc.add_table(len(software_list), 2)
        last_date = None
        for i, software in enumerate(software_list):
            left, right = tbl.cell(i, 0), tbl.cell(i, 1)
            left.width, right.width = self.date_col_width, self.item_col_width

            p = left.paragraphs[0]
            date = str(software['year'])
            if date != last_date:
                p.add_run(date)
            last_date = date

            p = right.paragraphs[0]
            name = add_hyperlink(p, software['name'], software['url'])
            name.font.bold = True

            descr = p.add_run(f": {software['description']}")
            descr.font.italic = True
        self._add_break(2)

    def _add_skills(self):
        self._new_section("SKILLS")
        skills_dict = self.data[self.CV_KEY]['skills']
        divs = 3
        for skill_key in skills_dict:
            self._new_subsection(skill_key.capitalize())
            skills = skills_dict[skill_key]
            skills.sort(key=lambda x: x['level'])
            tbl = self.doc.add_table(round(len(skills) / divs), divs)
            indent_table(tbl, 350)
            # tbl.style = "Light Grid Accent 1"
            for i, skill in enumerate(skills):
                row = i // divs
                col = i % divs
                cell = tbl.cell(row, col)
                p = cell.paragraphs[0]
                name = p.add_run(skill['name'])
                name.font.bold = True

                kw = p.add_run(f' ({", ".join(skill["keywords"])})')
                kw.font.italic = True
                p = cell.add_paragraph(" ")
                gap = Pt(5)
                p.runs[0].font.size = gap
                p.paragraph_format.line_spacing = gap
            self._add_break()
        self._add_break(2)

    def _add_works(self):
        self._new_section("LIST OF WORKS")
        works = self.data[self.WORKS_KEY]
        works.sort(key=lambda x: x['year'], reverse=True)
        last_date = None
        for i, work in enumerate(works):
            date = str(work['year'])
            if date != last_date:
                self._new_subsection(f"[ {date} ]")
            last_date = date
            p = self.doc.add_paragraph()
            name = p.add_run(work['name'])
            name.font.bold = True

            p.add_run(f" ({date}) ")
            subtitle = p.add_run(f"{work['subtitle']}. ")
            subtitle.font.italic = True

            p.add_run(f"{work['duration']}'")

            com = work['commission']
            if com:
                self._add_break(0.5)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = self.tab_size
                p.add_run(com)
            performances = work['performances']
            if performances:
                performances.sort(key=lambda x: x['date'], reverse=True)
                num_perf = len(performances)
                self._add_break(0.5)
                p = self.doc.add_paragraph("Performances")
                p.runs[0].font.italic = True
                p.paragraph_format.left_indent = self.tab_size
                self._add_break(0.5)
                for i, performance in enumerate(performances):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = self.tab_size * 2

                    event = p.add_run(performance['event'])
                    event.font.bold = True

                    if i == num_perf - 1:
                        p.add_run(" (world premiere)")

                    at = p.add_run(" @ ")
                    at.font.color.rgb = BLUE

                    _date = parse_date(performance['date'])
                    _date = f"{_date[1]} {_date[2]}, {_date[0]}."
                    p.add_run(
                        f"{performance['venue']}. {performance['city']}. {performance['country']}. {_date}")
                    for km in ['audio', 'video']:
                        url = performance.pop(km, None)
                        if url:
                            p.add_run(" ")
                            add_hyperlink(p, km, url)
                    self._add_break(0.5)


cv = CV()
cv.load_data(cv_path='cv.json', works_path='work-catalog.json')
cv.compile()
cv.write('cv.docx')
